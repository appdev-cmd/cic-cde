#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Append CDE CIC sections to Docs/01.Ban_HCM_Van_ban_de_xuat.docx
import re, sys, os
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

F = "Times New Roman"
CC = WD_ALIGN_PARAGRAPH.CENTER
JJ = WD_ALIGN_PARAGRAPH.JUSTIFY
LL = WD_ALIGN_PARAGRAPH.LEFT

# Read source markdown
md_path = os.path.join(os.path.dirname(__file__) or ".", "01.Ban_HCM_De_xuat_CDE_CIC.md")
with open(md_path, "r", encoding="utf-8") as fh:
    lines = fh.read().split("\n")

# Find where Section V starts (line with "## V. ĐỀ XUẤT NỀN TẢNG MÔI TRƯỜNG DỮ LIỆU CHUNG")
start_idx = 0
for idx, line in enumerate(lines):
    if line.strip().startswith("## V. ĐỀ XUẤT NỀN TẢNG MÔI TRƯỜNG DỮ LIỆU CHUNG"):
        start_idx = idx
        break

if start_idx == 0:
    # fallback if not found
    for idx, line in enumerate(lines):
        if line.strip().startswith("## V."):
            start_idx = idx
            break

print(f"Parsing CDE sections from line {start_idx + 1}: '{lines[start_idx]}'")

# Load the original docx document
doc_path = os.path.join(os.path.dirname(__file__) or ".", "01.Ban_HCM_Van_ban_de_xuat.docx")
doc = Document(doc_path)

# Remove the signature table at the end of the document
# The signature table is the last table in the document
if len(doc.tables) > 0:
    sig_table = doc.tables[-1]
    # Verify it is indeed the signature table
    # It has 1 row and 2 columns, and cell 1 contains "CÔNG TY CỔ PHẦN CÔNG NGHỆ"
    cell_txt = sig_table.cell(0, 1).text if len(sig_table.columns) > 1 else ""
    if "CÔNG TY" in cell_txt or "CIC" in cell_txt:
        print("Removing original signature table from docx...")
        p = sig_table._element.getparent()
        p.remove(sig_table._element)
    else:
        print("Warning: Last table does not look like the signature block. Not removing.")

# Styling helper functions
def sf(r, sz=13, b=False, it=False):
    r.font.name = F; r.font.size = Pt(sz); r.bold = b; r.italic = it
    rPr = r._element.get_or_add_rPr()
    rF = rPr.find(qn("w:rFonts"))
    if rF is None:
        rF = OxmlElement("w:rFonts"); rPr.insert(0, rF)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rF.set(qn(a), F)


def ap(text, sz=13, b=False, it=False, align=JJ, ind=True, bef=2, aft=2, li=None):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_before = Pt(bef)
    p.paragraph_format.space_after = Pt(aft)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    if ind:
        p.paragraph_format.first_line_indent = Pt(36)
    if li:
        p.paragraph_format.left_indent = Pt(li)
    
    # Process inline formatting like **bold** or *italic*
    parts = re.split(r"(\*\*.*?\*\*|\*[^*]+?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); sf(r, sz=sz, b=True, it=it)
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1]); sf(r, sz=sz, b=b, it=True)
        else:
            r = p.add_run(part); sf(r, sz=sz, b=b, it=it)
    return p


def bb(par):
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "1"); b.set(qn("w:color"), "000000")
    pbdr.append(b); pPr.append(pbdr)


def nbt(table):
    tblPr = table._tbl.tblPr; borders = OxmlElement("w:tblBorders")
    for s in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement("w:" + s); b.set(qn("w:val"), "none"); b.set(qn("w:sz"), "0")
        borders.append(b)
    tblPr.append(borders)


def sbt(table):
    tblPr = table._tbl.tblPr
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tb = OxmlElement("w:tblBorders")
    for s in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement("w:" + s); b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "000000"); tb.append(b)
    tblPr.append(tb)


def ct(cell, text, sz=11, b=False, it=False, align=LL):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.15
    
    parts = re.split(r"(\*\*.*?\*\*|\*[^*]+?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); sf(r, sz=sz, b=True, it=it)
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1]); sf(r, sz=sz, b=b, it=True)
        else:
            r = p.add_run(part); sf(r, sz=sz, b=b, it=it)


def clean(text):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    return text.strip()


# Spacing paragraph before appending CDE
ap("", sz=13, align=JJ, ind=False, bef=12, aft=12)

# ========== PARSING LOOP ==========
i = start_idx
in_table = False
table_rows = []

while i < len(lines):
    line = lines[i]
    s = line.strip()

    if not s or s == "---":
        i += 1; continue

    # Headers
    if s.startswith("## "):
        p = ap(s[3:].strip(), sz=13, b=True, align=LL, ind=False, bef=14, aft=4)
        i += 1; continue

    if s.startswith("### "):
        p = ap(s[4:].strip(), sz=13, b=True, align=LL, ind=False, bef=10, aft=3)
        i += 1; continue

    if s.startswith("#### "):
        p = ap(s[5:].strip(), sz=13, b=True, it=True, align=LL, ind=False, bef=8, aft=2)
        i += 1; continue

    # Table parsing
    if s.startswith("|") and not in_table:
        in_table = True; table_rows = [s]; i += 1; continue

    if in_table:
        if s.startswith("|"):
            table_rows.append(s); i += 1; continue
        else:
            in_table = False
            data = []
            for tr in table_rows:
                if "---" in tr or tr.strip() == "| | |" or tr.strip() == "|:---|:---|":
                    continue
                cells = [c.strip() for c in tr.split("|")[1:-1]]
                if cells:
                    data.append(cells)
            if len(data) >= 1:
                nc = len(data[0])
                tbl = doc.add_table(rows=len(data), cols=nc)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; tbl.autofit = True
                sbt(tbl)
                
                # Header row
                for j, val in enumerate(data[0]):
                    ct(tbl.rows[0].cells[j], val, sz=11, b=True, align=CC)
                
                # Data rows
                for r_idx, dr in enumerate(data[1:]):
                    row = tbl.rows[r_idx + 1]
                    for j, val in enumerate(dr[:nc]):
                        al = LL
                        if j == 0 and nc > 3: # STT col
                            al = CC
                        elif val.replace('.', '').replace(',', '').replace(' ', '').isdigit():
                            al = CC
                        elif "VNĐ" in val or "triệu" in val:
                            al = CC
                        ct(row.cells[j], val, sz=11, b=False, align=al)
            continue

    # Blockquotes
    if s.startswith("> "):
        p = doc.add_paragraph(); p.alignment = JJ
        p.paragraph_format.left_indent = Pt(36)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        
        txt = clean(s[2:])
        if txt.startswith("|"):
            pass
        else:
            r = p.add_run(txt)
            sf(r, sz=12, it=True)
        i += 1; continue

    # Bullet lists
    if s.startswith("- **"):
        text = s[2:].strip()
        m = re.match(r"^(\*\*.*?\*\*\s*:?\s*)(.*)", text)
        p = doc.add_paragraph(); p.alignment = JJ
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        
        if m:
            label = m.group(1)
            content_txt = m.group(2)
            r1 = p.add_run("• " + label)
            sf(r1, sz=13, b=True)
            
            parts = re.split(r"(\*\*.*?\*\*|\*[^*]+?\*)", content_txt)
            for part in parts:
                if not part: continue
                if part.startswith("**") and part.endswith("**"):
                    r2 = p.add_run(part[2:-2]); sf(r2, sz=13, b=True)
                elif part.startswith("*") and part.endswith("*"):
                    r2 = p.add_run(part[1:-1]); sf(r2, sz=13, it=True)
                else:
                    r2 = p.add_run(part); sf(r2, sz=13)
        else:
            r = p.add_run("• " + text)
            sf(r, sz=13)
        i += 1; continue

    if s.startswith("- "):
        p = doc.add_paragraph(); p.alignment = JJ
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run("• ")
        sf(r, sz=13)
        
        txt = s[2:].strip()
        parts = re.split(r"(\*\*.*?\*\*|\*[^*]+?\*)", txt)
        for part in parts:
            if not part: continue
            if part.startswith("**") and part.endswith("**"):
                r2 = p.add_run(part[2:-2]); sf(r2, sz=13, b=True)
            elif part.startswith("*") and part.endswith("*"):
                r2 = p.add_run(part[1:-1]); sf(r2, sz=13, it=True)
            else:
                r2 = p.add_run(part); sf(r2, sz=13)
        i += 1; continue

    # Numbered list
    if re.match(r"^(\d+[\.\)])\s+(.*)", s):
        m = re.match(r"^(\d+[\.\)])\s+(.*)", s)
        num_prefix = m.group(1)
        txt = m.group(2).strip()
        
        p = doc.add_paragraph(); p.alignment = JJ
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        
        r1 = p.add_run(num_prefix + " ")
        sf(r1, sz=13, b=True)
        
        parts = re.split(r"(\*\*.*?\*\*|\*[^*]+?\*)", txt)
        for part in parts:
            if not part: continue
            if part.startswith("**") and part.endswith("**"):
                r2 = p.add_run(part[2:-2]); sf(r2, sz=13, b=True)
            elif part.startswith("*") and part.endswith("*"):
                r2 = p.add_run(part[1:-1]); sf(r2, sz=13, it=True)
            else:
                r2 = p.add_run(part); sf(r2, sz=13)
        i += 1; continue

    # Plain paragraph
    if s:
        ap(s, sz=13, align=JJ, ind=True, bef=2, aft=2)
    i += 1

# ========== SIGNATURE TABLE ==========
doc.add_paragraph() # spacing
st = doc.add_table(rows=1, cols=2)
st.autofit = False; st.columns[0].width = Mm(80); st.columns[1].width = Mm(85); nbt(st)

st.cell(0, 0).paragraphs[0].clear()

right = st.cell(0, 1); right.paragraphs[0].clear()
p = right.paragraphs[0]; p.alignment = CC
sf(p.add_run("CÔNG TY CỔ PHẦN CÔNG NGHỆ VÀ TƯ VẤN CIC"), sz=11.5, b=True)
p2 = right.add_paragraph(); p2.alignment = CC
sf(p2.add_run("(Ký tên, đóng dấu)"), sz=11.5, it=True)

# Save document back to original path
doc.save(doc_path)
print(f"SUCCESS: Appended CDE to {doc_path}")

# Save copy to 01.Ban_HCM_Van_ban_de_xuat (1).docx as well
copy_path = os.path.join(os.path.dirname(__file__) or ".", "01.Ban_HCM_Van_ban_de_xuat (1).docx")
doc.save(copy_path)
print(f"SUCCESS: Copied to {copy_path}")
