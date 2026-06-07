#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Generate NĐ30 docx from the markdown file. All Vietnamese text is read from .md.
import re, sys, os
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

F = "Times New Roman"
CC = WD_ALIGN_PARAGRAPH.CENTER
JJ = WD_ALIGN_PARAGRAPH.JUSTIFY
LL = WD_ALIGN_PARAGRAPH.LEFT

# Read source markdown
md_path = os.path.join(os.path.dirname(__file__) or ".",
    "Đề xuất nhiệm vụ phát triển "
    "công nghệ chiến lược - CDE CIC.md")
with open(md_path, "r", encoding="utf-8") as fh:
    lines = fh.read().split("\n")

doc = Document()
sec = doc.sections[0]
sec.page_width = Mm(210); sec.page_height = Mm(297)
sec.top_margin = Mm(25); sec.bottom_margin = Mm(25)
sec.left_margin = Mm(30); sec.right_margin = Mm(20)
sec.different_first_page_header_footer = True

ns = doc.styles["Normal"]
ns.font.name = F; ns.font.size = Pt(13)
ns.element.rPr.rFonts.set(qn("w:eastAsia"), F)
pf = ns.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.15
pf.space_after = Pt(0); pf.space_before = Pt(0)


def sf(r, sz=13, b=False, it=False):
    r.font.name = F; r.font.size = Pt(sz); r.bold = b; r.italic = it
    rPr = r._element.get_or_add_rPr()
    rF = rPr.find(qn("w:rFonts"))
    if rF is None:
        rF = OxmlElement("w:rFonts"); rPr.insert(0, rF)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rF.set(qn(a), F)


def ap(text, sz=13, b=False, it=False, align=JJ, ind=True, bef=2, aft=2, li=None):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_before = Pt(bef)
    p.paragraph_format.space_after = Pt(aft)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    if ind:
        p.paragraph_format.first_line_indent = Pt(36)
    if li:
        p.paragraph_format.left_indent = Pt(li)
    r = p.add_run(text); sf(r, sz=sz, b=b, it=it)
    return p


def am(parts, align=JJ, ind=True, bef=2, aft=2, li=None):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_before = Pt(bef)
    p.paragraph_format.space_after = Pt(aft)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    if ind:
        p.paragraph_format.first_line_indent = Pt(36)
    if li:
        p.paragraph_format.left_indent = Pt(li)
    for text, sz, bold, italic in parts:
        r = p.add_run(text); sf(r, sz=sz, b=bold, it=italic)
    return p


def bb(par):
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "1"); b.set(qn("w:color"), "000000")
    pbdr.append(b); pPr.append(pbdr)


def nbt(table):
    tblPr = table._tbl.tblPr; borders = OxmlElement("w:tblBorders")
    for s in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement("w:" + s); b.set(qn("w:val"), "none"); b.set(qn("w:sz"), "0")
        borders.append(b)
    tblPr.append(borders)


def sbt(table):
    tblPr = table._tbl.tblPr
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tb = OxmlElement("w:tblBorders")
    for s in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement("w:" + s); b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "000000"); tb.append(b)
    tblPr.append(tb)


def ct(cell, text, sz=12.5, b=False, it=False, align=LL):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text); sf(r, sz=sz, b=b, it=it)


def clean(text):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    return text.strip()


# ========== HEADER ==========
ht = doc.add_table(rows=1, cols=2)
ht.alignment = WD_TABLE_ALIGNMENT.CENTER; ht.autofit = False
ht.columns[0].width = Mm(80); ht.columns[1].width = Mm(80); nbt(ht)

lc = ht.cell(0, 0); lc.paragraphs[0].clear()
# parse header from md lines 1-4
# left col: org name
left_org = clean(lines[2].split("|")[1]) if len(lines) > 2 else "CIC"
words = left_org.split()
mid = len(words) // 2
p = lc.paragraphs[0]; p.alignment = CC
sf(p.add_run(" ".join(words[:mid]).upper()), sz=12, b=True)
p2 = lc.add_paragraph(); p2.alignment = CC
sf(p2.add_run(" ".join(words[mid:]).upper()), sz=12, b=True); bb(p2)
p3 = lc.add_paragraph(); p3.alignment = CC
sf(p3.add_run("(CƠ QUAN, TỔ CHỨC, CÁ NHÂN ĐỀ XUẤT)"), sz=10, it=True)

rc = ht.cell(0, 1); rc.paragraphs[0].clear()
p = rc.paragraphs[0]; p.alignment = CC
sf(p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"), sz=12, b=True)
p2 = rc.add_paragraph(); p2.alignment = CC
sf(p2.add_run("Độc lập - Tự do - Hạnh phúc"), sz=13, b=True); bb(p2)
p3 = rc.add_paragraph(); p3.alignment = CC
sf(p3.add_run("......, ngày ... tháng ... năm ......"), sz=13, it=True)

# Title
ap("", sz=13, align=CC, ind=False, bef=10, aft=2)
title = clean(lines[5]) if len(lines) > 5 else "DE XUAT NHIEM VU"
ap(title, sz=14, b=True, align=CC, ind=False, bef=4, aft=10)

# ========== BODY: parse remaining lines ==========
i = 9  # skip processed header
in_table = False
table_rows = []

while i < len(lines):
    line = lines[i]
    s = line.strip()

    if not s or s == "---":
        i += 1; continue

    if s.startswith("<div") or s.startswith("</div"):
        i += 1; continue

    # Section ## headers
    if s.startswith("## "):
        ap(s[3:].strip(), sz=13, b=True, align=LL, ind=False, bef=8, aft=3)
        i += 1; continue

    # Sub-section ### headers
    if s.startswith("### "):
        ap(s[4:].strip(), sz=13, b=True, ind=False, bef=6, aft=2)
        i += 1; continue

    # Table detection
    if s.startswith("|") and not in_table:
        in_table = True; table_rows = [s]; i += 1; continue

    if in_table:
        if s.startswith("|"):
            table_rows.append(s); i += 1; continue
        else:
            in_table = False
            data = []
            for tr in table_rows:
                if "---" in tr:
                    continue
                cells = [c.strip() for c in tr.split("|")[1:-1]]
                if cells:
                    data.append(cells)
            if len(data) >= 2:
                nc = len(data[0])
                tbl = doc.add_table(rows=1, cols=nc)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; tbl.autofit = False
                if nc == 2:
                    tbl.columns[0].width = Mm(55); tbl.columns[1].width = Mm(105)
                elif nc == 3:
                    tbl.columns[0].width = Mm(26); tbl.columns[1].width = Mm(28); tbl.columns[2].width = Mm(106)
                sbt(tbl)
                for j, h in enumerate(data[0]):
                    ct(tbl.rows[0].cells[j], clean(h), sz=12.5, b=True, align=CC)
                for dr in data[1:]:
                    row = tbl.add_row()
                    for j, val in enumerate(dr):
                        al = LL if j == 0 else JJ
                        ct(row.cells[j], clean(val), sz=12.5, b=(j == 0 and nc >= 3), align=al)
            continue

    # Blockquote
    if s.startswith("> "):
        ap(clean(s[2:]), it=True, ind=False, bef=2, aft=2)
        i += 1; continue

    # Indented numbered (  1. 2.)
    if re.match(r"^\s+\d+\.", line):
        ap(clean(line.strip()), ind=False, bef=1, aft=1, li=30)
        i += 1; continue

    # Bullet with bold label: - **a) ...**
    if s.startswith("- **"):
        text = clean(s[2:])
        # Split at first :** or :**
        m2 = re.match(r"([^:]+:\s*)(.*)", text)
        if m2:
            am([(m2.group(1), 13, True, False), (m2.group(2), 13, False, False)],
               ind=False, bef=1, aft=1, li=18)
        else:
            ap(text, ind=False, bef=1, aft=1, li=18)
        i += 1; continue

    # Regular bullet
    if s.startswith("- "):
        ap("- " + clean(s[2:]), ind=False, bef=1, aft=1, li=18)
        i += 1; continue

    # Budget sub-items (indented with -)
    if re.match(r"^\s+- ", line):
        text = clean(line.strip()[2:])
        ap(text, ind=False, bef=0, aft=0, li=30)
        i += 1; continue

    # Bold item: **N. Title** or **N. Title: ** rest
    if s.startswith("**"):
        m3 = re.match(r"\*\*(.+?)\*\*\s*(.*)", s)
        if m3:
            label = m3.group(1); rest = clean(m3.group(2))
            if rest:
                # Check if rest starts with italic marker
                if rest.startswith("(") and ":" in rest:
                    am([(label + " ", 13, True, False), (rest, 13, False, True)],
                       ind=True, bef=3, aft=2)
                else:
                    am([(label, 13, True, False), (" " + rest, 13, False, False)],
                       ind=True, bef=3, aft=2)
            else:
                ap(label, b=True, ind=True, bef=3, aft=2)
        else:
            ap(clean(s), b=True, ind=True, bef=3, aft=2)
        i += 1; continue

    # Numbered top-level (1. 2. etc)
    if re.match(r"^\d+\.", s):
        ap(clean(s), ind=False, bef=1, aft=1, li=18)
        i += 1; continue

    # Plain paragraph
    text = clean(s)
    if text:
        ap(text, ind=True, bef=2, aft=2)
    i += 1

# ========== SIGNATURE ==========
ap("", ind=False, bef=8)
st = doc.add_table(rows=1, cols=2)
st.autofit = False; st.columns[0].width = Mm(80); st.columns[1].width = Mm(80); nbt(st)
st.cell(0, 0).paragraphs[0].clear()
right = st.cell(0, 1); right.paragraphs[0].clear()
p = right.paragraphs[0]; p.alignment = CC
sf(p.add_run("......, ngày ... tháng ... năm ......"), sz=13, it=True)
p2 = right.add_paragraph(); p2.alignment = CC
sf(p2.add_run("CƠ QUAN, TỔ CHỨC, CÁ NHÂN ĐỀ XUẤT"), sz=13, b=True)
p3 = right.add_paragraph(); p3.alignment = CC
sf(p3.add_run("(Ký tên, đóng dấu)"), sz=13, it=True)

# Page number from page 2
footer = sec.footer; footer.is_linked_to_previous = False
p = footer.paragraphs[0]; p.alignment = CC
run = p.add_run(); sf(run, sz=13)
f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
run._r.append(f1); run._r.append(instr); run._r.append(f2)
sec.first_page_footer.is_linked_to_previous = False

out = "De xuat nhiem vu CDE CIC (ND30).docx"
doc.save(out)
sys.stdout.buffer.write(("OK: " + out + " | " + str(len(doc.paragraphs)) + " paragraphs, " + str(len(doc.tables)) + " tables\n").encode("utf-8"))
