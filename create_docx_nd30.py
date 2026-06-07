# -*- coding: utf-8 -*-
"""Tạo DOCX chuyên nghiệp — Formal Navy palette + Cover Page — v2 fix."""
import re, os, sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ═══ FORMAL NAVY PALETTE ═══
C_PRIMARY = "1F3864"
C_ACCENT  = "C00000"
C_SECOND  = "595959"
C_BODY    = "202020"
C_BORDER  = "BFBFBF"
C_WHITE   = "FFFFFF"
C_ZEBRA   = "F2F2F2"
C_HDR_FILL = "1F3864"

FONT = "Times New Roman"
SZ_BODY = 13; SZ_H1 = 14; SZ_H2 = 13; SZ_H3 = 13
SZ_TBL = 10; SZ_SMALL = 11
SZ_COVER_TITLE = 22; SZ_COVER_SUB = 14
LINE_SP = Pt(19)
# Page content width in DXA (A4 - 3cm left - 2cm right = 16cm = 9072 DXA)
CONTENT_W = 9072

def sfont(run, sz=SZ_BODY, bold=False, italic=False, color=C_BODY, underline=False):
    run.font.name = FONT
    run.font.size = Pt(sz)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    if underline: run.underline = True
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>')
        rPr.insert(0, rFonts)
    else:
        for a in ['w:eastAsia','w:hAnsi','w:cs']:
            rFonts.set(qn(a), FONT)

def spara(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=Pt(4), before=Pt(0), line=None):
    p.alignment = align
    p.paragraph_format.space_after = after
    p.paragraph_format.space_before = before
    p.paragraph_format.line_spacing = line or LINE_SP

def add_runs(p, text, sz=SZ_BODY, bold=False, italic=False, color=C_BODY):
    # Clean markdown links: [text](url) → text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', text)
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); sfont(r, sz, bold=True, italic=italic, color=color)
        elif part.startswith('*') and part.endswith('*'):
            r = p.add_run(part[1:-1]); sfont(r, sz, bold=bold, italic=True, color=color)
        else:
            r = p.add_run(part); sfont(r, sz, bold=bold, italic=italic, color=color)

def add_bottom_border(p, color=C_ACCENT, sz=8):
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        p._element.insert(0, pPr)
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{sz}" w:space="3" w:color="{color}"/></w:pBdr>'))

def set_tbl_borders(table, color="000000", sz=4):
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None: tblPr.remove(old)
    bdr = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        + ''.join(f'<w:{e} w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
                  for e in ['top','left','bottom','right','insideH','insideV'])
        + '</w:tblBorders>')
    tblPr.append(bdr)

def shade_cell(cell, fill):
    tcPr = cell._element.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        cell._element.insert(0, tcPr)
    old = tcPr.find(qn('w:shd'))
    if old is not None: tcPr.remove(old)
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>'))

def set_cell_width(cell, width_dxa):
    """Set explicit cell width in DXA."""
    tcPr = cell._element.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        cell._element.insert(0, tcPr)
    old = tcPr.find(qn('w:tcW'))
    if old is not None: tcPr.remove(old)
    tcPr.append(parse_xml(f'<w:tcW {nsdecls("w")} w:type="dxa" w:w="{width_dxa}"/>'))

def clean_md(text):
    """Remove markdown formatting for length measurement."""
    t = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    t = re.sub(r'\*(.*?)\*', r'\1', t)
    t = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', t)
    return t.strip()

def calc_col_widths(headers, rows, total_w=CONTENT_W):
    """Calculate optimal column widths based on content."""
    cols = len(headers)
    # Measure max content length per column (in chars)
    max_lens = []
    for ci in range(cols):
        hdr_len = len(clean_md(headers[ci]))
        data_lens = []
        for row in rows:
            if ci < len(row):
                data_lens.append(len(clean_md(row[ci])))
        # Use avg of top values + header, not just max (to avoid one outlier)
        all_lens = sorted([hdr_len] + data_lens, reverse=True)
        # Representative length: avg of top 3 or all if fewer
        top = all_lens[:min(3, len(all_lens))]
        max_lens.append(max(sum(top) / len(top), hdr_len, 3))
    
    # Short column detection (like #, Năm, etc.)
    SHORT_THRESHOLD = 8
    short_cols = [i for i, l in enumerate(max_lens) if l <= SHORT_THRESHOLD]
    long_cols = [i for i in range(cols) if i not in short_cols]
    
    # Assign widths
    widths = [0] * cols
    
    # Short columns get fixed narrow width
    SHORT_W = 680  # ~1.2cm
    for ci in short_cols:
        # Scale slightly by content
        actual_chars = max_lens[ci]
        if actual_chars <= 3:
            widths[ci] = 500   # very narrow (#, Năm)
        elif actual_chars <= 5:
            widths[ci] = 650
        else:
            widths[ci] = 800
    
    used = sum(widths[ci] for ci in short_cols)
    remaining = total_w - used
    
    # Long columns share remaining proportionally
    if long_cols:
        total_long = sum(max_lens[ci] for ci in long_cols)
        if total_long > 0:
            for ci in long_cols:
                widths[ci] = int(remaining * max_lens[ci] / total_long)
    
    # Ensure minimum and adjust
    MIN_W = 600
    for ci in range(cols):
        if widths[ci] < MIN_W:
            widths[ci] = MIN_W
    
    # Normalize to total
    total = sum(widths)
    if total != total_w:
        ratio = total_w / total
        widths = [int(w * ratio) for w in widths]
        # Fix rounding
        diff = total_w - sum(widths)
        if diff != 0:
            widths[-1] += diff
    
    return widths

def add_md_table(doc, header_cells, data_rows):
    if not header_cells: return
    cols = len(header_cells)
    
    # Calculate widths
    col_widths = calc_col_widths(header_cells, data_rows)
    
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_tbl_borders(table, C_BORDER, 4)
    
    # Set table layout to fixed
    tblPr = table._element.find(qn('w:tblPr'))
    tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:type="dxa" w:w="{CONTENT_W}"/>'))
    tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))
    
    # Header row
    for ci, txt in enumerate(header_cells):
        cell = table.cell(0, ci)
        set_cell_width(cell, col_widths[ci])
        shade_cell(cell, C_HDR_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(14)
        clean = clean_md(txt)
        r = p.add_run(clean); sfont(r, SZ_TBL, bold=True, color=C_WHITE)
    
    # Data rows
    for ri, row in enumerate(data_rows):
        for ci in range(cols):
            cell = table.cell(ri + 1, ci)
            set_cell_width(cell, col_widths[ci])
            if ri % 2 == 1:
                shade_cell(cell, C_ZEBRA)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = Pt(13)
            
            txt = row[ci].strip() if ci < len(row) else ""
            # Clean links
            txt = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', txt)
            
            if '**' in txt:
                parts = re.split(r'(\*\*.*?\*\*)', txt)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = p.add_run(part[2:-2]); sfont(r, SZ_TBL, bold=True)
                    elif part:
                        pt = re.sub(r'\*(.*?)\*', r'\1', part)
                        r = p.add_run(pt); sfont(r, SZ_TBL)
            else:
                display = re.sub(r'\*(.*?)\*', r'\1', txt)
                r = p.add_run(display); sfont(r, SZ_TBL)
    
    # Spacing after
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run._element.append(parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>'))

def add_cover_page(doc):
    for _ in range(5):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    
    p = doc.add_paragraph()
    spara(p, WD_ALIGN_PARAGRAPH.CENTER, Pt(30), Pt(0))
    r = p.add_run("CÔNG TY CỔ PHẦN CÔNG NGHỆ VÀ TƯ VẤN CIC")
    sfont(r, 12, color=C_SECOND)

    p = doc.add_paragraph()
    spara(p, WD_ALIGN_PARAGRAPH.CENTER, Pt(20), Pt(10))
    add_bottom_border(p, C_ACCENT, 12)

    p = doc.add_paragraph()
    spara(p, WD_ALIGN_PARAGRAPH.CENTER, Pt(6), Pt(20))
    r = p.add_run("KẾ HOẠCH KINH DOANH 5 NĂM\n(2026 – 2030)")
    sfont(r, SZ_COVER_TITLE, bold=True, color=C_PRIMARY)

    p = doc.add_paragraph()
    spara(p, WD_ALIGN_PARAGRAPH.CENTER, Pt(6), Pt(8))
    r = p.add_run("Dự án Phát triển và Thương mại hóa\nNền tảng Môi trường Dữ liệu Chung BIM CDE CIC")
    sfont(r, SZ_COVER_SUB, italic=True, color=C_SECOND)

    p = doc.add_paragraph()
    spara(p, WD_ALIGN_PARAGRAPH.CENTER, Pt(40), Pt(20))
    add_bottom_border(p, C_ACCENT, 12)

    info = [
        ("Đơn vị chủ trì", "Công ty CP Công nghệ và Tư vấn CIC"),
        ("Sản phẩm", "Nền tảng CDE BIM – \"Make in Vietnam\""),
        ("Tổng vốn CAPEX", "25,1 tỷ VNĐ (CIC 70% + NATIF 30%)"),
        ("Doanh thu 5 năm", "252,6 tỷ VNĐ"),
        ("NPV / IRR", "+20,0 tỷ VNĐ / ~32%"),
        ("Hoàn vốn", "Quý I/2030 (~3,5 năm)"),
    ]
    tbl = doc.add_table(rows=len(info), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # No borders
    set_tbl_borders(tbl, C_WHITE, 0)
    tblPr = tbl._element.find(qn('w:tblPr'))
    tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:type="dxa" w:w="6800"/>'))
    tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))
    
    for i, (lbl, val) in enumerate(info):
        c0 = tbl.cell(i, 0); set_cell_width(c0, 2400)
        p = c0.paragraphs[0]; p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(lbl); sfont(r, 11, bold=True, color=C_PRIMARY)
        c1 = tbl.cell(i, 1); set_cell_width(c1, 4400)
        p = c1.paragraphs[0]; p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(val); sfont(r, 11, color=C_BODY)

    for _ in range(3):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    spara(p, WD_ALIGN_PARAGRAPH.CENTER, Pt(0), Pt(0))
    r = p.add_run("Tháng 06/2026"); sfont(r, 11, italic=True, color=C_SECOND)
    add_page_break(doc)

def add_page_number(doc):
    section = doc.sections[0]
    footer = section.footer; footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for tag in ['begin', None, 'end']:
        run = p.add_run(); sfont(run, 9, color=C_SECOND)
        if tag == 'begin':
            run._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        elif tag is None:
            run._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        else:
            run._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

def add_header_line(doc):
    section = doc.sections[0]
    header = section.header; header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_bottom_border(p, C_BORDER, 4)
    r = p.add_run("CIC Technology & Consulting  |  Kế hoạch kinh doanh 5 năm CDE CIC")
    sfont(r, 8, italic=True, color=C_BORDER)

# ═══════════════ MAIN PARSER ═══════════════
def convert(md_path, docx_path):
    print(f"📄 Đang xử lý: {os.path.basename(md_path)}")
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    
    add_cover_page(doc)
    add_page_number(doc)
    add_header_line(doc)
    
    lines = content.split('\n')
    i = 0
    in_table = False; t_hdr = []; t_sep = False; t_rows = []
    in_code = False; h1_count = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Code block
        if line.strip().startswith('```'):
            in_code = not in_code; i += 1; continue
        if in_code:
            p = doc.add_paragraph()
            spara(p, WD_ALIGN_PARAGRAPH.LEFT, Pt(0), Pt(0))
            r = p.add_run(line)
            r.font.name = "Consolas"; r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string(C_PRIMARY)
            i += 1; continue
        
        # Flush table
        if in_table and not line.strip().startswith('|'):
            if t_hdr: add_md_table(doc, t_hdr, t_rows)
            in_table = False; t_hdr = []; t_sep = False; t_rows = []
        
        # Skip empty / hr
        if not line.strip() or line.strip() in ['---','***','___']:
            i += 1; continue
        
        # ═══ HEADINGS ═══
        if line.startswith('#'):
            m = re.match(r'^(#{1,6})\s+(.*)', line)
            if m:
                lvl = len(m.group(1)); txt = clean_md(m.group(2))
                if lvl <= 2:
                    h1_count += 1
                    if h1_count <= 2: i += 1; continue
                
                if lvl == 2:
                    p = doc.add_paragraph()
                    spara(p, WD_ALIGN_PARAGRAPH.LEFT, Pt(8), Pt(22))
                    add_bottom_border(p, C_ACCENT, 8)
                    r = p.add_run(txt.upper()); sfont(r, SZ_H1, bold=True, color=C_ACCENT)
                elif lvl == 3:
                    p = doc.add_paragraph()
                    spara(p, WD_ALIGN_PARAGRAPH.LEFT, Pt(6), Pt(14))
                    r = p.add_run(txt); sfont(r, SZ_H2, bold=True, color=C_PRIMARY)
                elif lvl == 4:
                    p = doc.add_paragraph()
                    spara(p, WD_ALIGN_PARAGRAPH.LEFT, Pt(4), Pt(10))
                    r = p.add_run(txt); sfont(r, SZ_H3, bold=True, italic=True, color=C_BODY)
                else:
                    p = doc.add_paragraph()
                    spara(p, WD_ALIGN_PARAGRAPH.LEFT, Pt(3), Pt(8))
                    r = p.add_run(txt); sfont(r, SZ_BODY, bold=True, color=C_BODY)
                i += 1; continue
        
        # ═══ TABLE ═══
        if line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if not in_table:
                in_table = True; t_hdr = cells; t_sep = False; t_rows = []
            elif not t_sep:
                if all(c.strip().replace('-','').replace(':','') == '' for c in cells):
                    t_sep = True
                else: t_rows.append(cells)
            else: t_rows.append(cells)
            i += 1; continue
        
        # ═══ BULLET ═══
        if re.match(r'^[\s]*[-*]\s', line):
            indent = len(re.match(r'^(\s*)', line).group(1))
            text = re.sub(r'^[\s]*[-*]\s+', '', line).strip()
            # Clean markdown links
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            p = doc.add_paragraph()
            spara(p, WD_ALIGN_PARAGRAPH.JUSTIFY, Pt(2), Pt(2))
            lvl = indent // 2
            p.paragraph_format.left_indent = Cm(0.6 + lvl * 0.5)
            p.paragraph_format.first_line_indent = Cm(-0.35)
            r = p.add_run("• "); sfont(r, SZ_BODY, color=C_ACCENT)
            add_runs(p, text, SZ_BODY, color=C_BODY)
            i += 1; continue
        
        # ═══ NUMBERED LIST ═══
        if re.match(r'^[\s]*\d+[\.\)]\s', line):
            nm = re.match(r'^[\s]*(\d+[\.\)])\s+(.*)', line)
            if nm:
                indent = len(re.match(r'^(\s*)', line).group(1))
                text = nm.group(2).strip()
                # Clean markdown links
                text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
                p = doc.add_paragraph()
                spara(p, WD_ALIGN_PARAGRAPH.JUSTIFY, Pt(2), Pt(2))
                lvl = indent // 2
                p.paragraph_format.left_indent = Cm(0.6 + lvl * 0.5)
                p.paragraph_format.first_line_indent = Cm(-0.35)
                r = p.add_run(f"{nm.group(1)} "); sfont(r, SZ_BODY, bold=True, color=C_PRIMARY)
                add_runs(p, text, SZ_BODY, color=C_BODY)
            i += 1; continue
        
        # ═══ BLOCKQUOTE ═══
        if line.strip().startswith('>'):
            text = re.sub(r'^>\s*', '', line).strip()
            if text:
                text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
                p = doc.add_paragraph()
                spara(p, WD_ALIGN_PARAGRAPH.JUSTIFY, Pt(4), Pt(4))
                p.paragraph_format.left_indent = Cm(1)
                pPr = p._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
                    p._element.insert(0, pPr)
                pPr.append(parse_xml(
                    f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="16" w:space="8" w:color="{C_PRIMARY}"/></w:pBdr>'))
                add_runs(p, text, SZ_BODY, italic=True, color=C_SECOND)
            i += 1; continue
        
        # ═══ PARAGRAPH ═══
        text = line.strip()
        if text:
            if text.startswith('<') or text.startswith('</'): i += 1; continue
            # Clean links
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            p = doc.add_paragraph()
            spara(p, WD_ALIGN_PARAGRAPH.JUSTIFY, Pt(4), Pt(0))
            add_runs(p, text, SZ_BODY, color=C_BODY)
        i += 1
    
    if in_table and t_hdr: add_md_table(doc, t_hdr, t_rows)
    
    doc.save(docx_path)
    sz = os.path.getsize(docx_path)
    print(f"  ✅ Đã tạo: {os.path.basename(docx_path)} ({sz:,} bytes)")

if __name__ == "__main__":
    base = r"d:\01_Projects\cic-cde"
    convert(
        os.path.join(base, "Kế hoạch kinh doanh 5 năm CDE CIC.md"),
        os.path.join(base, "Kế hoạch kinh doanh 5 năm CDE CIC.docx")
    )
    print("\n🎉 Hoàn thành!")
